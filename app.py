import os
import io
import json
import re
import time
import http.cookiejar
import urllib.request
import urllib.parse
from datetime import datetime

import anthropic
import google.generativeai as genai
import pandas as pd
import yfinance as yf
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv
from flask import Flask, render_template, request, send_file

load_dotenv()

app = Flask(__name__)


@app.after_request
def _no_cache_html(response):
    if response.mimetype == "text/html":
        response.headers["Cache-Control"] = "no-store, no-cache, must-revalidate, max-age=0"
        response.headers["Pragma"] = "no-cache"
        response.headers["Expires"] = "0"
    return response


ANTHROPIC_API_KEY = os.getenv("ANTHROPIC_API_KEY")
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")

ANTHROPIC_MODELS = {
    "claude-opus-4-6",
    "claude-sonnet-4-6",
    "claude-haiku-4-5-20251001",
}
GEMINI_MODELS = {
    "gemini-2.5-flash",
    "gemini-2.5-flash-lite",
}
ALLOWED_MODELS = ANTHROPIC_MODELS | GEMINI_MODELS
DEFAULT_MODEL = "claude-sonnet-4-6"


# ---------------------------------------------------------------------------
# Yahoo Finance direct API helpers
# ---------------------------------------------------------------------------

_yahoo_crumb = None
_yahoo_opener = None


def _yahoo_session():
    """Create (or reuse) an HTTP opener with Yahoo Finance cookies + crumb."""
    global _yahoo_crumb, _yahoo_opener

    if _yahoo_crumb and _yahoo_opener:
        return _yahoo_opener, _yahoo_crumb

    cj = http.cookiejar.CookieJar()
    opener = urllib.request.build_opener(urllib.request.HTTPCookieProcessor(cj))
    opener.addheaders = [
        ("User-Agent",
         "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
         "(KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36"),
    ]

    # 1. Hit fc.yahoo.com to set cookies (it returns a 404, that's expected)
    try:
        opener.open("https://fc.yahoo.com", timeout=10)
    except Exception:
        pass  # 404 is normal – we only need the Set-Cookie header

    # 2. Fetch the crumb
    crumb_url = "https://query2.finance.yahoo.com/v1/test/getcrumb"
    resp = opener.open(crumb_url, timeout=10)
    crumb = resp.read().decode("utf-8")

    _yahoo_crumb = crumb
    _yahoo_opener = opener
    return opener, crumb


def _fetch_income_stmt_direct(ticker_symbol):
    """Fetch annual income statements directly from Yahoo Finance quoteSummary API.

    Returns a pandas DataFrame with financial items as rows and fiscal-year
    end dates as columns, matching the layout that yfinance produces.
    """
    opener, crumb = _yahoo_session()

    url = (
        f"https://query2.finance.yahoo.com/v10/finance/quoteSummary/"
        f"{urllib.parse.quote(ticker_symbol, safe='')}"
        f"?modules=incomeStatementHistory"
        f"&crumb={urllib.parse.quote(crumb, safe='')}"
    )
    resp = opener.open(url, timeout=15)
    data = json.loads(resp.read().decode("utf-8"))

    statements = (
        data["quoteSummary"]["result"][0]
        ["incomeStatementHistory"]["incomeStatementHistory"]
    )

    records = {}
    for stmt in statements:
        date_str = stmt["endDate"]["fmt"]           # e.g. "2024-06-30"
        col_key = pd.Timestamp(date_str)
        record = {}
        for key, val in stmt.items():
            if isinstance(val, dict) and "raw" in val:
                record[key] = val["raw"]
        records[col_key] = record

    df = pd.DataFrame(records)
    # Sort columns newest-first, like yfinance does
    df = df[sorted(df.columns, reverse=True)]
    return df


def _fetch_company_info_direct(ticker_symbol):
    """Fetch company profile from Yahoo Finance quoteSummary API."""
    try:
        opener, crumb = _yahoo_session()
        url = (
            f"https://query2.finance.yahoo.com/v10/finance/quoteSummary/"
            f"{urllib.parse.quote(ticker_symbol, safe='')}"
            f"?modules=quoteType,summaryProfile,price"
            f"&crumb={urllib.parse.quote(crumb, safe='')}"
        )
        resp = opener.open(url, timeout=15)
        data = json.loads(resp.read().decode("utf-8"))
        result = data["quoteSummary"]["result"][0]

        price_info = result.get("price", {})
        profile = result.get("summaryProfile", {})
        quote_type = result.get("quoteType", {})

        return {
            "longName": price_info.get("longName") or quote_type.get("longName"),
            "shortName": price_info.get("shortName") or quote_type.get("shortName"),
            "sector": profile.get("sector"),
            "industry": profile.get("industry"),
            "currency": price_info.get("currency"),
        }
    except Exception:
        return {}


# ---------------------------------------------------------------------------
# Formatting helpers
# ---------------------------------------------------------------------------

def format_number(value):
    """Format a number in millions."""
    if value is None:
        return {"formatted": "N/D", "negative": False, "pct": None}
    millions = value / 1_000_000
    negative = millions < 0
    formatted = f"{millions:,.1f} M"
    return {"formatted": formatted, "negative": negative, "pct": None}


def format_pct(value, revenue):
    """Format a number with its percentage over revenue."""
    if value is None or revenue is None or revenue == 0:
        return {"formatted": "N/D", "negative": False, "pct": None}
    millions = value / 1_000_000
    negative = millions < 0
    pct = (value / revenue) * 100
    formatted = f"{millions:,.1f} M"
    pct_str = f"({pct:+.1f}%)"
    return {"formatted": formatted, "negative": negative, "pct": pct_str}


def safe_get(df, key, col):
    """Safely get a value from a DataFrame."""
    try:
        return float(df.loc[key, col])
    except (KeyError, TypeError, ValueError):
        return None


# ---------------------------------------------------------------------------
# Main financial-data pipeline
# ---------------------------------------------------------------------------

# Maps our display labels to every possible index name across yfinance
# (pretty / camelCase) and the direct quoteSummary JSON keys.
KEY_MAP = {
    "Total Revenue": [
        "Total Revenue", "TotalRevenue", "totalRevenue",
    ],
    "Cost Of Revenue": [
        "Cost Of Revenue", "CostOfRevenue", "costOfRevenue",
    ],
    "Gross Profit": [
        "Gross Profit", "GrossProfit", "grossProfit",
    ],
    "Operating Expense": [
        "Operating Expense", "OperatingExpense", "totalOperatingExpenses",
        "Total Operating Expenses", "TotalExpenses", "operatingExpense",
    ],
    "Operating Income": [
        "Operating Income", "OperatingIncome", "operatingIncome",
    ],
    "EBITDA": [
        "EBITDA", "Normalized EBITDA", "NormalizedEBITDA", "ebitda",
    ],
    "Net Income": [
        "Net Income", "NetIncome", "netIncome",
        "Net Income Common Stockholders", "NetIncomeCommonStockholders",
        "netIncomeApplicableToCommonShares",
    ],
    "EBIT": [
        "EBIT", "ebit",
    ],
    "Interest Expense": [
        "Interest Expense", "InterestExpense", "interestExpense",
    ],
    "Tax Provision": [
        "Tax Provision", "TaxProvision", "incomeTaxExpense",
        "Income Tax Expense", "IncomeTaxExpense",
    ],
}


def get_financial_data(ticker_symbol):
    """Download and process financial data from Yahoo Finance."""

    income_stmt = None

    # --- Attempt 1: yfinance library (fast if session is healthy) ----------
    try:
        ticker = yf.Ticker(ticker_symbol)
        for attr in ("income_stmt", "financials"):
            try:
                stmt = getattr(ticker, attr)
                if stmt is not None and not stmt.empty:
                    income_stmt = stmt
                    break
            except Exception:
                pass
    except Exception:
        pass

    # --- Attempt 2: direct Yahoo Finance API call --------------------------
    if income_stmt is None:
        try:
            income_stmt = _fetch_income_stmt_direct(ticker_symbol)
        except Exception as exc:
            raise RuntimeError(
                f"No se pudieron obtener datos financieros para '{ticker_symbol}'. "
                f"Verifica que el ticker sea correcto y que haya conexion "
                f"a Internet. ({exc})"
            )

    if income_stmt is None or income_stmt.empty:
        raise RuntimeError(
            f"No se pudieron obtener datos financieros para '{ticker_symbol}'. "
            "Verifica que el ticker sea correcto."
        )

    # --- Company info (best-effort) ----------------------------------------
    info = {}
    try:
        info = getattr(ticker, "info", None) or {}
    except Exception:
        pass
    if not info.get("longName"):
        info = {**info, **_fetch_company_info_direct(ticker_symbol)}

    company_name = (
        info.get("longName") or info.get("shortName") or ticker_symbol.upper()
    )
    company_info = {
        "sector": info.get("sector", "N/D") or "N/D",
        "industry": info.get("industry", "N/D") or "N/D",
        "currency": info.get("currency", "USD") or "USD",
    }

    # --- Build table rows ---------------------------------------------------
    columns = income_stmt.columns[:4]
    years = [
        col.strftime("%Y") if hasattr(col, "strftime") else str(col)
        for col in columns
    ]

    def find_key(label):
        for candidate in KEY_MAP.get(label, [label]):
            if candidate in income_stmt.index:
                return candidate
        return None

    col_meta = []
    for col in columns:
        revenue = safe_get(income_stmt, find_key("Total Revenue"), col)
        col_meta.append({"col": col, "revenue": revenue})

    def build_row(label, key_label, css_class="", show_pct=True):
        key = find_key(key_label)
        cells = []
        for cm in col_meta:
            val = safe_get(income_stmt, key, cm["col"]) if key else None
            if show_pct:
                cells.append(format_pct(val, cm["revenue"]))
            else:
                cells.append(format_number(val))
        return {"label": label, "cells": cells, "css_class": css_class}

    financial_data = [
        build_row("Ingresos totales", "Total Revenue",
                  css_class="subtotal", show_pct=False),
        build_row("Coste de ventas", "Cost Of Revenue"),
        build_row("Margen bruto", "Gross Profit", css_class="subtotal"),
        build_row("Gastos operativos", "Operating Expense"),
        build_row("Resultado operativo (EBIT)", "Operating Income",
                  css_class="subtotal"),
        build_row("EBITDA", "EBITDA"),
        build_row("Gastos financieros", "Interest Expense"),
        build_row("Impuestos", "Tax Provision"),
        build_row("Beneficio neto", "Net Income", css_class="total"),
    ]

    raw_data = {}
    for year_label, col in zip(years, columns):
        year_data = {}
        for concept, key_label in [
            ("ingresos", "Total Revenue"),
            ("coste_ventas", "Cost Of Revenue"),
            ("margen_bruto", "Gross Profit"),
            ("gastos_operativos", "Operating Expense"),
            ("resultado_operativo", "Operating Income"),
            ("ebitda", "EBITDA"),
            ("gastos_financieros", "Interest Expense"),
            ("impuestos", "Tax Provision"),
            ("beneficio_neto", "Net Income"),
        ]:
            key = find_key(key_label)
            year_data[concept] = safe_get(income_stmt, key, col) if key else None
        raw_data[year_label] = year_data

    return financial_data, years, company_name, company_info, raw_data


# ---------------------------------------------------------------------------
# Claude memo generation
# ---------------------------------------------------------------------------

def _build_memo_prompt(company_name, ticker_symbol, years, raw_data, currency):
    data_text = ""
    for year in years:
        d = raw_data[year]
        data_text += f"\n--- {year} ---\n"
        for concept, value in d.items():
            if value is not None:
                data_text += (
                    f"  {concept}: {value / 1_000_000:,.1f} millones {currency}\n"
                )
            else:
                data_text += f"  {concept}: No disponible\n"

    return f"""Eres un analista financiero profesional. Redacta una nota de memoria financiera en español
con lenguaje formal contable sobre la empresa {company_name} (ticker: {ticker_symbol}).

Utiliza los siguientes datos reales de la cuenta de pérdidas y ganancias:

{data_text}

La nota debe incluir:
1. Un encabezamiento formal indicando que es la nota explicativa de la cuenta de pérdidas y ganancias
2. Un análisis de la evolución de los ingresos y el margen bruto
3. Un análisis de los gastos operativos y el resultado operativo
4. Un comentario sobre el EBITDA y su evolución
5. Un análisis del beneficio neto y la rentabilidad
6. Una conclusión con valoración general de la salud financiera

Usa formato profesional con párrafos bien estructurados. Incluye cifras concretas y porcentajes
de variación interanual. Redacta en un tono formal adecuado para una memoria anual corporativa.
No uses formato markdown. Escribe en texto plano con párrafos separados por líneas en blanco."""


def _generate_with_anthropic(prompt, model):
    if not ANTHROPIC_API_KEY:
        raise RuntimeError(
            "No se ha configurado la API key de Anthropic. "
            "Crea un archivo .env con ANTHROPIC_API_KEY=tu_clave"
        )
    client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
    message = client.messages.create(
        model=model,
        max_tokens=4096,
        messages=[{"role": "user", "content": prompt}],
    )
    return message.content[0].text


def _generate_with_gemini(prompt, model):
    if not GOOGLE_API_KEY:
        raise RuntimeError(
            "No se ha configurado la API key de Google. "
            "Crea un archivo .env con GOOGLE_API_KEY=tu_clave"
        )
    genai.configure(api_key=GOOGLE_API_KEY)
    gemini_model = genai.GenerativeModel(model)
    response = gemini_model.generate_content(
        prompt,
        generation_config={"max_output_tokens": 4096},
    )
    return response.text


def generate_memo(company_name, ticker_symbol, years, raw_data, currency, model):
    """Generate the financial memo using the selected AI provider."""
    if model not in ALLOWED_MODELS:
        raise RuntimeError(f"Modelo no soportado: {model}")

    prompt = _build_memo_prompt(
        company_name, ticker_symbol, years, raw_data, currency
    )

    if model in ANTHROPIC_MODELS:
        return _generate_with_anthropic(prompt, model)
    return _generate_with_gemini(prompt, model)


# ---------------------------------------------------------------------------
# Word document generation
# ---------------------------------------------------------------------------

def create_word_document(company_name, ticker_symbol, memo_text,
                         financial_data, years):
    """Create a Word document with the financial memo."""
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x1A, 0x23, 0x32)

    title = doc.add_heading(level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Nota de Memoria Financiera")
    run.font.color.rgb = RGBColor(0x0A, 0x24, 0x63)
    run.font.size = Pt(22)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"{company_name} ({ticker_symbol.upper()})")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1E, 0x52, 0x99)
    run.bold = True

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run(
        f"Fecha de elaboracion: {datetime.now().strftime('%d/%m/%Y')}"
    )
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x71, 0x80, 0x96)

    doc.add_paragraph("")

    heading = doc.add_heading("Datos Financieros", level=1)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x0A, 0x24, 0x63)

    num_cols = 1 + len(years)
    table = doc.add_table(rows=1, cols=num_cols)
    table.style = "Light Grid Accent 1"

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Concepto"
    for i, year in enumerate(years):
        hdr_cells[i + 1].text = year

    for row_data in financial_data:
        row_cells = table.add_row().cells
        row_cells[0].text = row_data["label"]
        for i, val in enumerate(row_data["cells"]):
            text = val["formatted"]
            if val["pct"] is not None:
                text += f" {val['pct']}"
            row_cells[i + 1].text = text

    doc.add_paragraph("")

    heading = doc.add_heading("Nota Explicativa", level=1)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x0A, 0x24, 0x63)

    for paragraph_text in memo_text.split("\n\n"):
        paragraph_text = paragraph_text.strip()
        if paragraph_text:
            doc.add_paragraph(paragraph_text)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ---------------------------------------------------------------------------
# Flask routes
# ---------------------------------------------------------------------------

@app.route("/", methods=["GET", "POST"])
def index():
    context = {
        "current_year": datetime.now().year,
        "selected_model": DEFAULT_MODEL,
    }

    if request.method == "POST":
        ticker_symbol = request.form.get("ticker", "").strip().upper()
        model = request.form.get("model", DEFAULT_MODEL).strip()
        if model not in ALLOWED_MODELS:
            model = DEFAULT_MODEL
        context["ticker"] = ticker_symbol
        context["selected_model"] = model

        if not ticker_symbol:
            context["error"] = "Por favor, introduce un ticker valido."
            return render_template("index.html", **context)

        try:
            financial_data, years, company_name, company_info, raw_data = (
                get_financial_data(ticker_symbol)
            )
            currency = company_info.get("currency", "USD")

            memo_raw = generate_memo(
                company_name, ticker_symbol, years, raw_data, currency, model
            )
            memo_html = memo_raw.replace("\n\n", "</p><p>").replace("\n", "<br>")
            memo_html = f"<p>{memo_html}</p>"

            app.config[f"memo_{ticker_symbol}"] = memo_raw
            app.config[f"data_{ticker_symbol}"] = financial_data
            app.config[f"years_{ticker_symbol}"] = years
            app.config[f"name_{ticker_symbol}"] = company_name

            context.update({
                "financial_data": financial_data,
                "years": years,
                "company_name": company_name,
                "company_info": company_info,
                "memo": memo_html,
            })

        except Exception as e:
            context["error"] = f"Error al procesar los datos: {str(e)}"

    return render_template("index.html", **context)


@app.route("/download/<ticker>")
def download(ticker):
    ticker = ticker.upper()
    memo = app.config.get(f"memo_{ticker}")
    financial_data = app.config.get(f"data_{ticker}")
    years = app.config.get(f"years_{ticker}")
    company_name = app.config.get(f"name_{ticker}")

    if not memo:
        return "No hay datos disponibles. Realiza primero el analisis.", 404

    buffer = create_word_document(
        company_name, ticker, memo, financial_data, years
    )
    filename = f"Nota_Memoria_{company_name.replace(' ', '_')}_{ticker}.docx"

    return send_file(
        buffer,
        as_attachment=True,
        download_name=filename,
        mimetype=(
            "application/vnd.openxmlformats-officedocument"
            ".wordprocessingml.document"
        ),
    )


if __name__ == "__main__":
    app.run(debug=True, host="0.0.0.0", port=5000)
